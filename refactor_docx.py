import docx


def refactor_docx(path, data, filename):
    doc = docx.Document(path)
    # work with text
    for paragraph in doc.paragraphs:  # get all paragraphs
        for key in data.keys():  # compare keys with words in paragraph
            if key in paragraph.text:  # if key in paragraph we change it
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = docx.shared.Pt(10)
                paragraph.text = paragraph.text.replace(key, data.get(key))

    # work with tables
    for table in doc.tables:  # get all tables
        for row in table.rows:  # get all rows in table
            for cell in row.cells:  # get all cells in row
                for paragraph in cell.paragraphs:  # get all paragraphs in cell
                    for key in data.keys():  # compare keys with words in paragraph
                        if key in paragraph.text:  # if key is in paragraph we change it

                            style = doc.styles['Normal']
                            font = style.font
                            font.name = 'Times New Roman'
                            font.size = docx.shared.Pt(10)
                            paragraph.text = paragraph.text.replace(key, data.get(key))

    doc.save(f'{filename}.docx')


if __name__ == '__main__':
    refactor_docx(path, data, filename)
